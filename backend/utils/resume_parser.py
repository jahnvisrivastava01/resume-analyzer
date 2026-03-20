import os
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_file(file_path):

    file_extension = os.path.splitext(file_path)[1].lower()

    # PDF
    if file_extension == ".pdf":
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    # DOCX
    elif file_extension == ".docx":
        doc = Document(file_path)
        text = ""

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"

    # Extract tables properly
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text += cell.text + "\n"

    # Extract headers and footers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if para.text:
                text += para.text + "\n"

        footer = section.footer
        for para in footer.paragraphs:
            if para.text:
                text += para.text + "\n"

    return text

